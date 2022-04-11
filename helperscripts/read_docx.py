from __future__ import annotations
from unicodedata import category
import docx
from docx.text.paragraph import Paragraph
from docx.document import Document
import pandas as pd

import glob
import os
import re
import pickle
from typing import Tuple, List, Dict
from dataclasses import dataclass, field

from featureextraction import format_extractor

@dataclass
class EssayStructured:
    # raw_doc: Document
    file_name: str
    student_id: str
    title: List[Paragraph] = field(default_factory=list)
    body: List[Paragraph] = field(default_factory=list)
    reference: List[Paragraph] = field(default_factory=list)
    others: List[Paragraph] = field(default_factory=list)
    scores: Dict[str, int] = field(default_factory=dict)

    def preview(self,):
        """This is for debugging purposes"""
        print(f"\n>>>TITLE")
        for i in self.title:
            print(i.text)

        print(f"\n>>>BODY")
        for i in self.body:
            print(i.text[:80], "...")
        
        print(f"\n>>>REFERENCE")
        for i in self.reference:
            print(i.text)

    def get_title(self): 
        return '\n'.join([i.text.strip() for i in self.title])

    def get_body(self): 
        return '\n'.join([i.text.strip() for i in self.body])
    
    def get_reference(self): 
        return '\n'.join([i.text.strip() for i in self.reference])

    def get_word_count(self):
        return sum([len(i.text.split(" ")) for i in self.body])

    def to_list(self) -> List[str]:
        return [self.file_name, self.student_id, *[self.para_to_text(i) for i in [self.title, self.body, self.reference]]]

    def para_to_text(self, para: List[Paragraph]) -> str:
        return '\n'.join(p.text for p in para)

    def attach_score(self, scores: Dict[str, int]):
        self.scores = scores

@dataclass
class DocumentBin():
    hw_code: str
    semester_code: str
    essays: List[EssayStructured]
    
    def attach_scores(self, scores: pd.DataFrame):
        categories = ['content', 'research', 'communication', 'organization', 'bibliography', 'efforts', 'quality of writing']
        for essay in self.essays:
            grades = {j: scores.loc[essay.student_id, j] for j in categories}
            essay.attach_score(grades)
        

    def to_csv(self, path: str, encoding: str='utf-8'):
        array2d = [i.to_list() for i in self.essays]
        df = pd.DataFrame(array2d, columns=['file_name', 'netID', 'title', 'body', 'reference'])
        df.to_csv(path, encoding=encoding)
    
    def save_to_disk(self, path: str):
        *_, file_name = os.path.split(path)
        has_semester_code = re.search(r'(?:fa|sp|su)\d{2}', file_name)
        has_hw_code = re.search(r"hw\d", file_name)
        if not has_semester_code:
            raise Exception(f"file name does not have semester code ({file_name})")
        if not has_hw_code:
            raise Exception(f"file name does not have hw code ({file_name})")
        
        essays = self.essays
        with open(path, 'wb') as f:
            pickle.dump(essays, f)

    @staticmethod
    def load_from_saved(path: str) -> DocumentBin:
        *_, file_name = os.path.split(path)
        hw_code, semester_code, *_ = file_name.split('_')
        with open(path, 'rb') as f:
            return DocumentBin(semester_code, hw_code, pickle.load(f))

class DocxReader():
    def __init__(self, path_to_hw: str=None, *args, **kargs):
        self.path_to_essays = path_to_hw

    def get_netID_to_structured_content(self) -> Dict[str, EssayStructured]:
        result = {}
        for i in self.get_docx_assignment_list(self.path_to_essays):
            netID = self.get_netID_from_file_path(i)
            structured_essay = self.read_docx_and_structure(i)
            # structured_essay = self.parse_document_into_sections(i)

            result[netID] = structured_essay

        return result

    def get_netID_from_file_path(self, file_path:str) -> str:
        netID_pattern = r'\w{2}\d{6}'
        match_result = re.search(netID_pattern, file_path)
        if match_result == None:
            result = ''
        else:
            result = match_result.group()    
        return result
    
    def read_document(self, file_path: str) -> Document:
        """This method return the raw xml format of the docx file"""
        _, ext = os.path.splitext(file_path)
        assert ext == ".docx"
        
        document = docx.Document(file_path)
        
        return document

    def parse_document_into_sections(self, path: str) -> EssayStructured:
        document = self.read_document(path)
        file_name = os.path.split(path)[-1]
        student_id = re.search(r'\w{2}\d{6}', file_name).group()
        paragraphs = document.paragraphs
        essayStructured = EssayStructured(file_name, student_id)
        
        found_content_start = False
        found_ref_start = False

        for para in paragraphs:
            word_count = len(para.text.split(' '))
            is_empty_paragraph = re.match(r'^\s*$',para.text) != None
            if is_empty_paragraph: continue # skip empty paragraphs

            if not found_content_start:
                if word_count >= 45 or para.text.lower() in ['abstract']:
                    found_content_start = True
                    essayStructured.body.append(para)
                else:
                    essayStructured.title.append(para)
            elif not found_ref_start:
                if word_count <= 3:
                    is_page_number = re.match(r'.*\d\s?$', para.text) != None # page number e.g. "page 1" or "Bobby 1"
                    is_other_title = para.text.lower() in ['introduction', 'conclusion', 'abstract']
                    if is_page_number or is_other_title:
                        continue
                
                    found_ref_start = True
                    essayStructured.reference.append(para)
                else:
                    essayStructured.body.append(para)
            else:
                essayStructured.reference.append(para)
        
        return essayStructured

    def read_docx_and_structure(self, file_path: str) -> tuple:
        """
        This method will disect an essay as 'title page', 'body', and 'bibliography' and return them
        """
        head, ext = os.path.splitext(file_path)
        assert ext == '.docx'

        document = docx.Document(file_path)
        # effort_feature_vectors = self.extract_effort_feature(document)
        effort_feature_vectors = []
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip() != '']
        
        content_start_index: int
        reference_start_index: int

        found_content_start = False
        found_reference_start = False
        if len(paragraphs) == 0:
            return ('','','')

        for i, paragraph in enumerate(paragraphs):
            word_count = len(paragraph.split(' '))
            
            if (not found_content_start):
                if word_count >= 45 or paragraph.lower() in ['abstract']: 
                    found_content_start = True
                    content_start_index = i
            elif not found_reference_start:
                if word_count<=3:
                    # if (not found_reference_start) and paragraph.lower() in ['bibliography', 'references', 'reference','sources']:
                    is_page_number = re.match(r'.*\d\s?$', paragraph) != None # page number e.g. "page 1" or "Bobby 1"
                    is_other_title = paragraph.lower() in ['introduction', 'conclusion', 'abstract']
                    if is_page_number or is_other_title:
                        continue
                
                    found_reference_start = True
                    reference_start_index = i

        if found_content_start:
            title = '\n'.join(paragraphs[:content_start_index])
            
            if found_reference_start:
                body = '\n'.join(paragraphs[content_start_index: reference_start_index])
                references = '\n'.join(paragraphs[reference_start_index:])
            else:
                body = '\n'.join(paragraphs[content_start_index:])
                references = ''
        else:
            title = ''
            body = '\n'.join(paragraphs)
            references = ''
            
        return (title, body, references, *effort_feature_vectors)

    @staticmethod
    def get_docx_assignment_list(path_to_essay_folder: str) -> List[str]:
        result = glob.glob(f"{path_to_essay_folder}/*.docx")
        result = list(filter(lambda x: '~$' not in x,result))
        return result

    def extract_effort_feature(self, document: docx.document.Document) -> Tuple:
        """ return (line space portion, justification portion, indentation portion, font portion)
        """
        line_spaces, justifications, indentations, font_portion= [], [], [], []
        # target_values= {
        #     "line_space": 1.15,
        #     "justitification": None,
        #     "indentation": 0.5,
        #     "font": "Times New Roman"
        # }

        for i in document.paragraphs:
            pf = i.paragraph_format
            first_chararacter_is_tab = i.text[0] == "\t" if len(i.text) > 0 else False
            
            line_spaces.append(True if pf.line_spacing == 1.15 else False)
            justifications.append(True if pf.alignment == None else False)
            
            if pf.first_line_indent:
                indentations.append(True if pf.first_line_indent.inches == 0.5 else False)
            else:
                indentations.append(first_chararacter_is_tab)
            
            font_portion.append(True if format_extractor.get_majority_font_name(i)=="Times New Roman" else False)
            
        result = tuple( get_portion_of_true(i) for i in (line_spaces, justifications, indentations, font_portion) )
        return result

    
def get_portion_of_true(array: List[bool]) -> float:
    trues = sum([1 for i in array if i is True])
    return trues/len(array)


if __name__=="__main__":
    DIR = './essays/hw2_sp22'
    docx_name = "Assignment 2. Privacy_ac487426_attempt_2022-03-09-03-41-20_Privacy.docx"

    reader = DocxReader(DIR)
    # print(os.getcwd())
    document = docx.Document(f"{DIR}/{docx_name}")
    features = reader.extract_effort_feature(document)
    print(f"{features}")
    # test_pdf = 'Assignment 1. Normative Ethics_ec529919_attempt_2020-09-15-06-45-50_EvonChoong_300Z_Assignment1.pdf'
    # first_file_path = get_docx_assignment_list(DIR)[0]
    # content = read_docx_content(first_file_path)
    # print(f'\n{get_netID(first_file_path)}')
    # print(len(''.join(content)))
    # read_pdf_content(f"{DIR}/{test_pdf}")
    # print(read_docx_content('./testing.docx').encode('unicode_escape'))
