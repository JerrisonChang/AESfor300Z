from __future__ import annotations
import docx
from docx.document import Document

import glob
import os
import re
from typing import Tuple, List, Dict

from helperscripts.featureextraction import format_extractor
from .essaystructured import EssayStructured

class DocxReader():
    def __init__(self, path_to_hw: str=None, *args, **kargs):
        self.path_to_essays = path_to_hw

    def get_netID_to_structured_content(self) -> Dict[str, EssayStructured]:
        result = {}
        for i in self.get_docx_assignment_list(self.path_to_essays):
            netID = self.get_netID_from_file_path(i)
            structured_essay = self.read_docx_and_structure(i)
            # structured_essay = self.parse_document_into_sections(i)

            result[netID] = structured_essay

        return result

    def get_netID_from_file_path(self, file_path:str) -> str:
        netID_pattern = r'\w{2}\d{6}'
        match_result = re.search(netID_pattern, file_path)
        if match_result == None:
            result = ''
        else:
            result = match_result.group()    
        return result
    
    def read_document(self, file_path: str) -> Document:
        """This method return the raw xml format of the docx file"""
        _, ext = os.path.splitext(file_path)
        assert ext == ".docx"
        
        document = docx.Document(file_path)
        
        return document

    def parse_document_into_sections(self, path: str) -> EssayStructured:
        document = self.read_document(path)
        file_name = os.path.split(path)[-1]
        student_id = re.search(r'\w{2}\d{6}', file_name).group()
        paragraphs = document.paragraphs
        essayStructured = EssayStructured(file_name, student_id)
        
        found_content_start = False
        found_ref_start = False

        for para in paragraphs:
            word_count = len(para.text.split(' '))
            is_empty_paragraph = re.match(r'^\s*$',para.text) != None
            if is_empty_paragraph: continue # skip empty paragraphs

            if not found_content_start:
                if word_count >= 45 or para.text.lower() in ['abstract']:
                    found_content_start = True
                    essayStructured.body.append(para)
                else:
                    essayStructured.title.append(para)
            elif not found_ref_start:
                if word_count <= 3:
                    is_page_number = re.match(r'.*\d\s?$', para.text) != None # page number e.g. "page 1" or "Bobby 1"
                    is_other_title = para.text.lower() in ['introduction', 'conclusion', 'abstract']
                    if is_page_number or is_other_title:
                        continue
                
                    found_ref_start = True
                    essayStructured.reference.append(para)
                else:
                    essayStructured.body.append(para)
            else:
                essayStructured.reference.append(para)
        
        return essayStructured

    def read_docx_and_structure(self, file_path: str) -> tuple:
        """
        This method will disect an essay as 'title page', 'body', and 'bibliography' and return them
        """
        _, ext = os.path.splitext(file_path)
        assert ext == '.docx'

        document = docx.Document(file_path)
        effort_feature_vectors = []
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip() != '']
        
        content_start_index: int
        reference_start_index: int

        found_content_start = False
        found_reference_start = False
        if len(paragraphs) == 0:
            return ('','','')

        for i, paragraph in enumerate(paragraphs):
            word_count = len(paragraph.split(' '))
            
            if (not found_content_start):
                if word_count >= 45 or paragraph.lower() in ['abstract']: 
                    found_content_start = True
                    content_start_index = i
            elif not found_reference_start:
                if word_count<=3:
                    # if (not found_reference_start) and paragraph.lower() in ['bibliography', 'references', 'reference','sources']:
                    is_page_number = re.match(r'.*\d\s?$', paragraph) != None # page number e.g. "page 1" or "Bobby 1"
                    is_other_title = paragraph.lower() in ['introduction', 'conclusion', 'abstract']
                    if is_page_number or is_other_title:
                        continue
                
                    found_reference_start = True
                    reference_start_index = i

        if found_content_start:
            title = '\n'.join(paragraphs[:content_start_index])
            
            if found_reference_start:
                body = '\n'.join(paragraphs[content_start_index: reference_start_index])
                references = '\n'.join(paragraphs[reference_start_index:])
            else:
                body = '\n'.join(paragraphs[content_start_index:])
                references = ''
        else:
            title = ''
            body = '\n'.join(paragraphs)
            references = ''
            
        return (title, body, references, *effort_feature_vectors)

    @staticmethod
    def get_docx_assignment_list(path_to_essay_folder: str) -> List[str]:
        result = glob.glob(f"{path_to_essay_folder}/*.docx")
        result = list(filter(lambda x: '~$' not in x,result))
        return result

    def extract_effort_feature(self, document: docx.document.Document) -> Tuple:
        """ return (line space portion, justification portion, indentation portion, font portion)
        """
        line_spaces, justifications, indentations, font_portion= [], [], [], []
        # target_values= {
        #     "line_space": 1.15,
        #     "justitification": None,
        #     "indentation": 0.5,
        #     "font": "Times New Roman"
        # }

        for i in document.paragraphs:
            pf = i.paragraph_format
            first_chararacter_is_tab = i.text[0] == "\t" if len(i.text) > 0 else False
            
            line_spaces.append(True if pf.line_spacing == 1.15 else False)
            justifications.append(True if pf.alignment == None else False)
            
            if pf.first_line_indent:
                indentations.append(True if pf.first_line_indent.inches == 0.5 else False)
            else:
                indentations.append(first_chararacter_is_tab)
            
            font_portion.append(True if format_extractor.get_majority_font_name(i)=="Times New Roman" else False)
            
        result = tuple( get_portion_of_true(i) for i in (line_spaces, justifications, indentations, font_portion) )
        return result

    
def get_portion_of_true(array: List[bool]) -> float:
    trues = sum([1 for i in array if i is True])
    return trues/len(array)
